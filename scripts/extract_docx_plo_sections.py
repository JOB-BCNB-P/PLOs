from pathlib import Path
from docx import Document
import re

sources = [
    Path('/home/ubuntu/upload/หลักสูตรพยบ.2570แก้ไข16ก.ค.2569.docx'),
]
for source in sources:
    document = Document(source)
    lines = []
    for paragraph in document.paragraphs:
        text = ' '.join(paragraph.text.split())
        if text and re.search(r'PLO|ผลลัพธ์การเรียนรู้|1\.1|2\.1|3\.1', text, re.IGNORECASE):
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            text = ' | '.join(' '.join(cell.text.split()) for cell in row.cells)
            if text and re.search(r'PLO|ผลลัพธ์การเรียนรู้|1\.1|2\.1|3\.1', text, re.IGNORECASE):
                lines.append(text)
    out = Path('/home/ubuntu/plo-assessment-system/docs/source_analysis/docx_plo_extract.txt')
    out.write_text('\n'.join(lines), encoding='utf-8')
    print(f'{source.name}: matched_lines={len(lines)} output={out}')
